import argparse
import csv
import configparser
import datetime as dt

import requests
import re
import io

from pathlib import Path
from collections import defaultdict
from datetime import datetime
from docx import Document
from docx.enum.text import WD_BREAK

DAY_RE = re.compile(r"\.\s+\d{2}\.\d{2}\.")
TIME_RE = re.compile(r"^\d{2}:\d{2}$")
LF_RE = re.compile(r"\bLF\s*\d+(?:\.\d+)?\b")
DATE_RE = re.compile(r"\d{2}\.\d{2}\.")

ALIASES = {
    "LF 1.1": "LF 1.1 – IT-Unternehmen",
    "LF 1.2": "LF 2.1 – Arbeitsmethoden",
    "LF 1.4": "LF 1.4 – Präsentationstechniken",
    "LF 2.4": "LF 2.4 – Hard- und Software",
    "LF 5.2": "LF 5.2 – Algorithmen und Diagramme",
    "Englisch IT": "Englisch – Fachenglisch",
    "GSL": "GSL",
    "Bewegungstherapie": "Bewegungstherapie",
}

def read_static_config(path="config.ini") -> dict:
    cfg = ConfigParser()
    cfg.read(path, encoding="utf-8")
    return {
        "NAME": cfg.get("user", "name", fallback=""),
        "YEAR": cfg.get("user", "year", fallback=""),
        "CLASS": cfg.get("user", "class", fallback=""),
    }

def insert_pagebreak_before_week2(doc: Document):
    for p in doc.paragraphs:
        if "{{%2_START%}}" in p.text:
            pb = p.insert_paragraph_before()
            pb.add_run().add_break(WD_BREAK.PAGE)
            return

def normalize_key(entry: str) -> str:
    entry = entry.strip()

    # 1️⃣ Lernfeld extrahieren
    m = LF_RE.search(entry)
    if m:
        # Leerzeichen normieren → "LF 2.4"
        return re.sub(r"\s+", " ", m.group(0)).strip()

    # 2️⃣ sonstige Einheiten (Lehrername entfernen)
    for sep in (" - ", " | "):
        if sep in entry:
            return entry.split(sep, 1)[0].strip()

    return entry



def duration_hours(start_str, end_str):
    fmt = "%H:%M"
    start = datetime.strptime(start_str, fmt)
    end   = datetime.strptime(end_str, fmt)
    return (end - start).total_seconds() / 3600


def is_slot_row(row):
    if len(row) < 4:
        return False
    return bool(
        TIME_RE.match((row[2] or "").strip()) and
        TIME_RE.match((row[3] or "").strip())
    )

    return

def is_week_header(row):
    if len(row) < 9:
        return False

    cells = row[4:9]
    return all(DAY_RE.search((c or "").strip()) for c in cells)


def monday_of_week(date: dt.date) -> dt.date:
    return date - dt.timedelta(days=date.weekday())

def parse_args():
    p = argparse.ArgumentParser()
    p.add_argument("--config", required=True)
    p.add_argument("--template", required=True)
    p.add_argument("--out", required=True)
    return p.parse_args()

def load_config(path: Path):
    cfg = configparser.ConfigParser()
    read_files = cfg.read(path, encoding="utf-8")

    if not read_files:
        raise FileNotFoundError(f"Config-Datei nicht gefunden/lesbar: {path.resolve()}")

    if "berichtsheft" not in cfg:
        raise KeyError(
            f"Sektion [berichtsheft] fehlt in {path.resolve()}.\n"
            f"Gefundene Sektionen: {cfg.sections()}"
        )

    sec = cfg["berichtsheft"]

    # optional: eigene Sektion für Personendaten, sonst fallback auf [berichtsheft]
    person = cfg["person"] if "person" in cfg else sec

    return {
        "sheet_csv_url": sec.get("sheet_csv_url", "").strip(),
        "week_mode": sec.get("week_mode", "current").strip(),
        "manual_week_start": sec.get("manual_week_start", "").strip(),
        "NAME": person.get("name", "").strip(),
        "YEAR": person.get("year", "").strip(),
        "CLASS": person.get("class", "").strip(),
    }

def build_days_for_two_weeks(schedule, week_start: dt.date):
    ordered = []
    for offset in range(0, 5):  # Woche 1
        d = week_start + dt.timedelta(days=offset)
        key = d.strftime("%d.%m.")  # muss zu deinen schedule-Keys passen
        subjects = schedule.get(key, defaultdict(float))
        total = sum(subjects.values())
        ordered.append((key, total, subjects))

    for offset in range(7, 12):  # Woche 2
        d = week_start + dt.timedelta(days=offset)
        key = d.strftime("%d.%m.")
        subjects = schedule.get(key, defaultdict(float))
        total = sum(subjects.values())
        ordered.append((key, total, subjects))

    return ordered

def format_day(date_str: str, subjects: dict, total: float):
    # Datum als Header (so wie du es unter dem Wochentag haben willst)
    header = date_str  # z.B. "23.02."

    # SLOT → Themenliste (gern mit Aliases)
    items = sorted(subjects.items(), key=lambda x: (-x[1], x[0]))
    slot_lines = [f"- {ALIASES.get(k, k)} ({h:g}h)" for k, h in items]
    slot = "\n".join(slot_lines)

    hour = f"{total:g}"
    return header, slot, hour

def build_content(days, cfg):
    content = {"NAME": cfg["NAME"], "YEAR": cfg["YEAR"], "CLASS": cfg["CLASS"]}

    for i, (date_str, total, subjects) in enumerate(days[:10]):
        header, slot, hour = format_day(date_str, subjects, total)
        content[f"{i}_HEADER"] = header
        content[f"{i}_SLOT"] = slot
        content[f"{i}_HOUR"] = hour

    return content

def fill_template(content, template_path: Path, out_path: Path):
    import re

    doc = Document(str(template_path))
    insert_pagebreak_before_week2(doc)

    def replace(text):
        for key, value in content.items():
            pattern = r"\{\{%\s*" + re.escape(key) + r"\s*%\}\}"
            text = re.sub(pattern, str(value), text)
        return text

    # normale Paragraphen
    for p in doc.paragraphs:
        p.text = replace(p.text)

    # Tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = replace(p.text)


    doc.save(str(out_path))




def resolve_week_start(cfg) -> dt.date:
    today = dt.date.today()
    if cfg["week_mode"] == "next":
        return monday_of_week(today) + dt.timedelta(days=7)
    if cfg["week_mode"] == "manual" and cfg["manual_week_start"]:
        return dt.date.fromisoformat(cfg["manual_week_start"])
    return monday_of_week(today)

def fetch_csv(url: str) -> str:
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    return r.text
    
def main():

    schedule: dict[str, dict[str, float]] = defaultdict(lambda: defaultdict(float))
    args = parse_args()
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    cfg = load_config(Path(args.config))

    week_start = resolve_week_start(cfg)
    week_end = week_start + dt.timedelta(days=4)

    csv_text = fetch_csv(cfg["sheet_csv_url"])  
    

    csv_reader = csv.reader(io.StringIO(csv_text), delimiter=",")

    for row in csv_reader:

        if is_week_header(row):
            current_days = current_days = [re.findall(r"\d{2}\.\d{2}\.", cell)[0] for cell in row[4:9]]



        if is_slot_row(row):
            dur = duration_hours(row[2].strip(), row[3].strip())

            for day, entry in zip(current_days, row[4:9]):
                entry = (entry or "").strip()
                if not entry or entry.upper() == "FREI":
                    continue

                key = normalize_key(entry)
                schedule[day][key] += dur

    days = build_days_for_two_weeks(schedule, week_start)
    content = build_content(days, cfg)

    def fmt(d: dt.date) -> str:
        return d.strftime("%d.%m.")

    content["1_START"] = fmt(week_start)
    content["1_END"]   = fmt(week_start + dt.timedelta(days=4))
    content["2_START"] = fmt(week_start + dt.timedelta(days=7))
    content["2_END"]   = fmt(week_start + dt.timedelta(days=11))

    out_file = out_dir / "Ausbildungsnachweis_Wochenblatt_KW.docx"

    fill_template(content, Path(args.template), out_file)

if __name__ == "__main__":
    main()